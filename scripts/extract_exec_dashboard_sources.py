import json
import re
import sys
import zipfile
from pathlib import Path

import openpyxl
from docx import Document


SOURCE_FILES = [
    Path(r"C:\Users\jhhea\Downloads\Reforecast 5-26 Working Model.xlsx"),
    Path(r"C:\Users\jhhea\Downloads\UHP Cash Flow May-Dec - 5-29 working model.xlsx"),
    Path(r"C:\Users\jhhea\Downloads\Consolidated_Grounds_Crew_Equipment_Needs.docx"),
    Path(r"C:\Users\jhhea\Downloads\UHP_Tech_Transformation_Deck_v2.pptx"),
]


def clean(value):
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def interesting_row(values):
    text = " ".join(clean(v).lower() for v in values)
    terms = [
        "capex",
        "capital",
        "opex",
        "op ex",
        "operating",
        "revenue",
        "tuition",
        "cash",
        "expense",
        "grounds",
        "equipment",
        "technology",
        "tech",
        "may",
        "dec",
        "budget",
        "forecast",
    ]
    return any(term in text for term in terms)


def summarize_workbook(path):
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    result = {"file": path.name, "sheets": []}
    for ws in wb.worksheets:
        rows = []
        max_rows = min(ws.max_row or 0, 120)
        max_cols = min(ws.max_column or 0, 18)
        for row in ws.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
            if any(cell is not None and clean(cell) for cell in row) and (len(rows) < 25 or interesting_row(row)):
                rows.append([clean(cell) for cell in row])
            if len(rows) >= 60:
                break
        result["sheets"].append({
            "name": ws.title,
            "size": [ws.max_row, ws.max_column],
            "rows": rows,
        })
    return result


def summarize_docx(path):
    doc = Document(path)
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    tables = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows[:80]:
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                table_rows.append(values)
        tables.append(table_rows)
    equipment = []
    for table in tables:
        if not table or "Equipment" not in " ".join(table[0]):
            continue
        for row in table[1:]:
            if len(row) < 6 or not row[0].isdigit():
                continue
            cost_raw = row[2]
            normalized = cost_raw.replace(",", "").lower().strip()
            parsed_cost = None
            multi_k = re.match(r"(\d+)\s*-\s*\$?(\d+)k", normalized)
            if multi_k:
                parsed_cost = int(multi_k.group(1)) * int(multi_k.group(2)) * 1000
            elif "v1" not in normalized:
                nums = re.findall(r"\$?([0-9]+(?:\.[0-9]+)?)", normalized)
                if nums:
                    parsed_cost = float(nums[0]) * (3 if "x3" in normalized else 1)
            equipment.append({
                "item": row[1].split(" | ")[0],
                "costRaw": cost_raw,
                "cost": parsed_cost,
                "priority": row[3],
                "decision": row[5],
            })
    known = [item for item in equipment if item["cost"] is not None]
    go = [item for item in known if item["decision"].lower() in {"go", "used"}]
    top = [item for item in known if item["priority"].lower() == "top"]
    return {
        "file": path.name,
        "paragraphs": paragraphs[:80],
        "tables": tables,
        "equipmentSummary": {
            "knownCostTotal": sum(item["cost"] for item in known),
            "goOrUsedTotal": sum(item["cost"] for item in go),
            "topPriorityKnownTotal": sum(item["cost"] for item in top),
            "knownCostCount": len(known),
            "goOrUsedCount": len(go),
        },
        "equipment": equipment,
    }


def summarize_pptx(path):
    slides = []
    with zipfile.ZipFile(path) as zf:
        names = sorted(
            [n for n in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)],
            key=lambda n: int(re.search(r"slide(\d+)\.xml", n).group(1)),
        )
        for name in names:
            xml = zf.read(name).decode("utf-8", errors="ignore")
            texts = re.findall(r"<a:t>(.*?)</a:t>", xml)
            decoded = [clean(re.sub(r"&amp;", "&", t)) for t in texts if clean(t)]
            slides.append({"slide": len(slides) + 1, "text": decoded})
    return {"file": path.name, "slides": slides}


def main():
    output = []
    for path in SOURCE_FILES:
        if not path.exists():
            output.append({"file": str(path), "error": "missing"})
            continue
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            output.append(summarize_workbook(path))
        elif suffix == ".docx":
            output.append(summarize_docx(path))
        elif suffix == ".pptx":
            output.append(summarize_pptx(path))
    json.dump(output, sys.stdout, indent=2)


if __name__ == "__main__":
    main()
